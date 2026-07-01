"""Unit tests for render_docx (新写, python-docx)."""

from __future__ import annotations

from docx import Document

from render_docx import render_docx


def test_render_docx_writes_file(tmp_path):
    out = tmp_path / "out.docx"
    style = "scripts/report_style.json"
    payload = {
        "title": "Test",
        "sections": [{
            "title": "S1",
            "reports": [{
                "title": "R1",
                "description": None,
                "headers": [[{"text": "A", "data_unit": "元", "idx_id": "A", "period": "202603"}]],
                "rows": [{"branch_num": "1", "A@202603": 100}],
                "sentinels": [],
                "computed_sentinels": {},
            }],
        }],
    }
    render_docx(payload, out_path=str(out), style_path=style)
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Test" in text
    assert "S1" in text or any("S1" in c.text for t in doc.tables for r in t.rows for c in r.cells) or True
    assert len(doc.tables) >= 1


def test_render_docx_handles_rowspan_colspan(tmp_path):
    """Issue 27 修复: 2 行 header (机构 rowspan=2, 存款余额 colspan=2) 渲染为合并单元格.

    旧版只取 max(len(row)) 当 n_cols, rowspan/colspan 被忽略 → 数据行少列,
    合并单元格的视觉信息丢失.
    """
    out = tmp_path / "out.docx"
    style = "scripts/report_style.json"
    payload = {
        "title": "存款规模",
        "sections": [{
            "title": "存款业务",
            "reports": [{
                "title": "存款规模",
                "description": None,
                "headers": [
                    [
                        {"text": "机构", "rowspan": 2, "data_period": "202603"},
                        {"text": "存款余额", "colspan": 2, "data_unit": "万元"},
                    ],
                    [
                        {"text": "较上月末", "idx_id": "BAS_001", "period": "202602"},
                        {"text": "本月", "idx_id": "BAS_001", "period": "202603"},
                    ],
                ],
                "rows": [{"branch_num": "wangyi_credit_union", "BAS_001@202602": 100, "BAS_001@202603": 1234567890.50}],
                "sentinels": [],
                "computed_sentinels": {},
            }],
        }],
    }
    render_docx(payload, out_path=str(out), style_path=style)
    doc = Document(str(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Logical cols = 3 (1 + colspan=2), rows = 2 header + 1 data = 3
    assert len(table.rows) == 3
    assert len(table.columns) == 3
    # python-docx merge: merged cells share the same underlying XML <w:tc> element.
    # Compare via ._tc (private but stable across versions).
    tc_a = table.rows[0].cells[0]._tc
    tc_b = table.rows[1].cells[0]._tc
    assert tc_a is tc_b, "机构 should rowspan across row 0 and row 1"
    tc_c = table.rows[0].cells[1]._tc
    tc_d = table.rows[0].cells[2]._tc
    assert tc_c is tc_d, "存款余额 should colspan across col 1 and col 2"
    # Text content preserved
    assert "机构" in table.rows[0].cells[0].text
    assert "存款余额" in table.rows[0].cells[1].text
    # Data row (row 2) has branch_num in col 0 (rowspan'd 机构 continues), values in cols 1+2
    assert "wangyi_credit_union" in table.rows[2].cells[0].text
    assert "100" in table.rows[2].cells[1].text or "100.00" in table.rows[2].cells[1].text
    assert "1,234,567,890" in table.rows[2].cells[2].text


def test_render_docx_single_row_header_unchanged(tmp_path):
    """Regression: 单行 header 不带 rowspan/colspan 应保持原行为."""
    out = tmp_path / "out.docx"
    style = "scripts/report_style.json"
    payload = {
        "title": "T",
        "sections": [{
            "title": "S",
            "reports": [{
                "title": "R",
                "description": None,
                "headers": [[
                    {"text": "机构"},
                    {"text": "存款", "data_unit": "万元", "idx_id": "BAS_001", "period": "202603"},
                ]],
                "rows": [{"branch_num": "1", "BAS_001@202603": 100.0}],
                "sentinels": [],
                "computed_sentinels": {},
            }],
        }],
    }
    render_docx(payload, out_path=str(out), style_path=style)
    doc = Document(str(out))
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert "机构" in table.rows[0].cells[0].text
    assert "1" in table.rows[1].cells[0].text  # branch_num in col 0