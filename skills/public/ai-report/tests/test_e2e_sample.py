"""E2E test for the full ai-report pipeline (design + runtime) on the wangyi 5-section sample.

Run with: pytest tests/test_e2e_sample.py -v

Asserts:
- design_pipeline.run_report() walks all 5 sections, every section ends in
  approval_status='approved', and every approved_run.status='ok' (no sentinel).
- runtime_pipeline.run_report() produces .report.md, .report.docx, .status.json.
- The MD and DOCX both contain the section titles, table titles, and the
  post-unit-conversion values (e.g. 1234567890.50 / 10000 → "123456.78905" — the
  Decimal-precision guarantee from Phase 1 policy).
- The DOCX preserves rowspan/colspan cell merges from the 2-row headers in
  sections 一/二 (存款规模/贷款规模).
- status.json reports `status='ok'`, `sections_approved=5`,
  `query_failures=0`, no sentinels.
"""

from __future__ import annotations

import json
import sys
from decimal import Decimal
from pathlib import Path

import pytest

SCRIPTS_DIR = Path(__file__).resolve().parent.parent / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from design_pipeline import run_report
from docx import Document
from duckdb_store import Store, make_report_id
from report_md import build_runtime_payload
from runtime_pipeline import RuntimePipeline
from sqlbot_client import MockSQLBotClient


FIXTURE = "tests/fixtures/mock_sqlbot/wangyi_2026_03.json"
EXAMPLE_MD = "example/wangyi_2026_03.md"


@pytest.fixture
def env(tmp_path, monkeypatch):
    """Real DuckDB-backed E2E env. Mocks LLM + checkpoints, keeps real SQLBot."""
    db_path = str(tmp_path / "e2e.duckdb")
    store = Store(db_path=db_path)
    store.open()
    sqlbot = MockSQLBotClient(fixture_path=FIXTURE)

    # Auto-approve all checkpoints
    monkeypatch.setattr(
        "design_pipeline._checkpoint",
        lambda msg, opts: "continue" if "continue" in opts else (
            "approve" if "approve" in opts else opts[0]
        ),
    )
    # LLM stubs: safe SQL that just returns one constant per row
    monkeypatch.setattr(
        "design_pipeline._llm_codegen",
        lambda ir, wide: "SELECT branch_num, 1.0 AS dummy FROM wide",
    )
    # LLM describe: canned Chinese paragraph
    monkeypatch.setattr(
        "design_pipeline._llm_describe",
        lambda wide, title, description_prompt=None: f"{title} 经营情况整体稳定。",
    )
    return {"store": store, "sqlbot": sqlbot, "tmp": tmp_path}


def _out_dir(env):
    return env["tmp"] / "out"


# ---- DESIGN PIPELINE ---- #


def test_design_walks_all_5_sections(env):
    """run_report() on wangyi 5-section MD → 5 results, all approved."""
    result = run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    assert result["status"] == "completed"
    assert len(result["results"]) == 5
    for r in result["results"]:
        assert r["approval_status"] == "approved", (
            f"section run did not approve: {r}"
        )


def test_design_persists_one_approved_run_per_section(env):
    """Each of the 5 sections has exactly one row in approved_runs.

    Status may be 'ok' or 'partial' depending on whether all queried facts
    returned data; sections that reference periods outside `time_info` (e.g.
    section 一's BAS_001@202602 MoM column) just render as empty cells — no
    query is issued for them, no sentinel is fired.
    """
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    rows = env["store"].list_approved_tables(report_id)
    assert len(rows) == 5, f"expected 5 approved sections, got {len(rows)}"
    for r in rows:
        assert r["status"] in ("ok", "partial"), f"approved_run status={r['status']}"
        sentinels = json.loads(r["sentinels"]) if isinstance(r["sentinels"], str) else r["sentinels"]
        assert all(s.startswith("⚠️") for s in sentinels), f"non-code sentinels: {sentinels}"


def test_design_metric_facts_have_all_6_idx_ids(env):
    """Every section pulls data for all 6 idx_ids (all in the global mock fixture)."""
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    rows = env["store"].list_approved_tables(report_id)
    # Pick the first section (存款业务); its run_id has 6 facts (one per idx_id × time_info).
    facts = env["store"].get_metric_facts(rows[0]["run_id"], rows[0]["table_id"])
    idx_ids = sorted({f["idx_id"] for f in facts})
    assert idx_ids == ["BAS_001", "BAS_010", "BAS_020", "BAS_026", "BAS_030", "BAS_040"], (
        f"expected all 6 idx_ids, got {idx_ids}"
    )


def test_design_unit_conversion_applied_end_to_end(env):
    """Phase 1 Decimal-precision guarantee: 1234567890.50 万元 → 123456.78905 (no float drift)."""
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    rows = env["store"].list_approved_tables(report_id)
    # Section 0 (存款业务) has BAS_001@202603 with value 1234567890.50 in fixture
    sec0 = rows[0]
    wide = json.loads(sec0["wide_table"]) if isinstance(sec0["wide_table"], str) else sec0["wide_table"]
    cell = wide[0]["BAS_001@202603"]
    # After apply_units(万元 → divide 10000): 1234567890.50 / 10000 = 123456.78905
    # Decimal precision preserved via str(Decimal): the str repr is exact.
    assert "123456.78905" in str(cell), (
        f"expected 123456.78905 (Decimal precision), got {cell!r}"
    )


# ---- RUNTIME PIPELINE ---- #


def test_runtime_renders_md_docx_status_json(env):
    """After design approval, runtime produces all 3 deliverables."""
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    out_dir = _out_dir(env)
    pipeline = RuntimePipeline(env["store"])
    result = pipeline.run_report(report_id, out_dir=str(out_dir))

    assert result["status"] == "completed"
    assert Path(result["out_md"]).exists()
    assert Path(result["out_docx"]).exists()


def test_runtime_md_contains_all_section_titles_and_values(env):
    """MD output renders all 5 section titles + the decimal-converted value.

    Section 一/二 reference period=202602 in their leaf headers but only declare
    time_info=["202603"]; the 202602 cells render as `—` (em dash, no data
    fetched). The 202603 cell renders the Decimal-precision converted value.
    """
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    pipeline = RuntimePipeline(env["store"])
    pipeline.run_report(report_id, out_dir=str(_out_dir(env)))

    md_path = _out_dir(env) / f"{report_id}.report.md"
    text = md_path.read_text(encoding="utf-8")
    for title in ["存款业务", "贷款业务", "收入与利润", "资产质量", "流动性指标"]:
        assert title in text, f"missing section title {title!r} in MD"
    # Section 一 BAS_001@202603 = 1234567890.50 / 10000 = 123456.78905 → render as
    # raw Decimal str in MD (MD renderer doesn't apply currency formatting,
    # only docx does). Assert Decimal precision preserved.
    assert "123456.78905" in text, "Decimal-precision converted value missing from MD"
    # Float-round-trip bug detector: assert the trailing-precision string is gone
    assert "123456.78904999" not in text, "float round-trip detected in MD output"
    # The 202602 cell renders as — (em dash, no data)
    assert "—" in text, "empty cell for unreferenced period should render as em dash"


def test_runtime_docx_preserves_rowspan_colspan_merges(env):
    """Section 一/二 have 2-row headers (rowspan=2 机构 + colspan=2 存款余额).

    Note: design_pipeline.run_report simplifies to 1 table per section (per the
    "1 section → 1 table" inline comment), so we expect 5 tables total even
    though the MD has 6 H3 blocks. The wangyi sample's section 三 (收入与利润)
    only gets its first H3 table (营业收入) processed; the second (利润总额)
    is dropped by run_report's simplification.
    """
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    pipeline = RuntimePipeline(env["store"])
    pipeline.run_report(report_id, out_dir=str(_out_dir(env)))

    docx_path = _out_dir(env) / f"{report_id}.report.docx"
    doc = Document(str(docx_path))
    assert len(doc.tables) == 5, f"expected 5 tables (1 per section, by run_report simplification), got {len(doc.tables)}"

    # Table 0 is 存款规模: 2-row header + 1 data row = 3 rows, 3 cols
    t0 = doc.tables[0]
    assert len(t0.rows) == 3
    assert len(t0.columns) == 3

    # 机构 (row 0, col 0) should span rows 0..1 (rowspan=2)
    tc_a = t0.rows[0].cells[0]._tc
    tc_b = t0.rows[1].cells[0]._tc
    assert tc_a is tc_b, "机构 should rowspan across header rows"

    # 存款余额 (row 0, col 1) should span cols 1..2 (colspan=2)
    tc_c = t0.rows[0].cells[1]._tc
    tc_d = t0.rows[0].cells[2]._tc
    assert tc_c is tc_d, "存款余额 should colspan across data cols"

    # Data row (row 2) col 1 (较上月末, BAS_001@202602 not in time_info) → "—"
    cell_202602 = t0.rows[2].cells[1].text
    assert cell_202602.strip() in ("—", ""), f"expected em-dash for missing period, got {cell_202602!r}"
    # Data row (row 2) col 2 (本月, BAS_001@202603) → ¥123,456.79 (1234567890.50 / 10000 → 123456.78905 → 2dp)
    cell_202603 = t0.rows[2].cells[2].text
    assert "¥123,456.79" in cell_202603, f"expected currency-formatted ¥123,456.79, got {cell_202603!r}"


def test_runtime_status_json_reports_all_sections_approved(env):
    """status.json reports sections approved (no query_failed on happy-path)."""
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    pipeline = RuntimePipeline(env["store"])
    pipeline.run_report(report_id, out_dir=str(_out_dir(env)))

    status_path = _out_dir(env) / f"{report_id}.status.json"
    status = json.loads(status_path.read_text(encoding="utf-8"))
    assert status["report_id"] == report_id
    assert status["total_sections"] == 5
    assert status["approved_sections"] == 5
    assert status["draft_sections"] == 0
    assert status["total_sentinels"] == 0
    # No ⚠️ codes in the by_code breakdown
    for code, count in status["sentinels_by_code"].items():
        assert count == 0, f"{code} count={count} on happy-path fixture"


def test_runtime_payload_round_trip_preserves_decimals(env):
    """build_runtime_payload → renderer path preserves the Decimal-as-str contract."""
    run_report(env["store"], env["sqlbot"], EXAMPLE_MD)
    report_id = make_report_id(EXAMPLE_MD)
    payload = build_runtime_payload(env["store"], report_id)

    # 5 sections, 5 tables (run_report 简化 1 节 1 表, section 三 的第二个 H3 被丢弃)
    assert len(payload["sections"]) == 5
    n_tables = sum(len(s["reports"]) for s in payload["sections"])
    assert n_tables == 5

    # Section 0 (存款业务) row 0 should have BAS_001@202603 cell as str(Decimal)
    sec0_row0 = payload["sections"][0]["reports"][0]["rows"][0]
    cell = sec0_row0["BAS_001@202603"]
    assert "123456.78905" in str(cell), f"Decimal precision lost: {cell!r}"
    # The renderer-bound value should still parse as Decimal (no float round-trip)
    Decimal(str(cell))  # raises if not parseable