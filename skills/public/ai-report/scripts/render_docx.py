"""ai-report pure docx renderer (新写, python-docx).

Phase 1 fix (Issue 27): handle thead rowspan/colspan via cell merging.
python-docx doesn't auto-resolve rowspan/colspan attrs — _build_header_layout
pre-computes the logical grid + per-cell spans + per-col data source, then
_render_report merges and fills in a single pass.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Cm, Pt, RGBColor

DATA_TYPE_MAP = {"元": "currency", "万元": "currency", "亿元": "currency", "%": "percentage", "百分点": "ratio"}
_ORG_LABELS = {"机构", "行社", "分行", "网点", "org"}


def _load_style(path: str) -> dict:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _apply_font(run, font_cfg: dict) -> None:
    run.font.name = font_cfg.get("name", "宋体")
    run.font.size = Pt(font_cfg.get("size", 11))
    run.font.bold = bool(font_cfg.get("bold", False))
    if "color" in font_cfg:
        run.font.color.rgb = RGBColor.from_string(font_cfg["color"].lstrip("#"))


def _set_cell_text(cell, text: str, *, main_font: dict, sub_font: dict | None = None) -> None:
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    run = p.add_run(text)
    _apply_font(run, main_font)
    if sub_font:
        sub_p = cell.add_paragraph()
        sub_run = sub_p.add_run(sub_font["text"])
        _apply_font(sub_run, {**main_font, **sub_font})


def _format_value(value, data_type: str) -> str:
    if value in (None, "", "⚠️QUERY_FAILED", "⚠️COMPUTE_FAILED"):
        return str(value) if value else ""
    try:
        v = float(value)
    except (ValueError, TypeError):
        return str(value)
    if data_type == "percentage":
        return f"{v * 100:.1f}%"
    if data_type == "currency":
        return f"¥{v:,.2f}"
    if data_type == "ratio":
        return f"{v:.2f}"
    return f"{v:,.0f}"


def _as_dict(th: Any) -> dict:
    if isinstance(th, dict):
        return th
    return {"text": str(th)}


def _span(th: Any, attr: str) -> int:
    th = _as_dict(th)
    n = th.get(attr) or 1
    try:
        return max(1, int(n))
    except (ValueError, TypeError):
        return 1


def _build_header_layout(headers: list[list[Any]]) -> tuple[int, list[tuple[int, int, int, int, dict]], list[dict | None], list[str | None]]:
    """Compute everything _render_report needs in one pass.

    Returns:
      n_cols: logical column count (= max sum of colspan across header rows)
      spans: list of (anchor_r, anchor_c, rowspan, colspan, th_dict) for each cell,
             in header-iteration order. Used for merging.
      col_sources: per-data-col th_dict (or None). Leaf (last-row) wins; if no
                   leaf at col_idx, deepest rowspan cell from above wins.
                   None means no header covers that col (rendered as empty).
      col_units: per-data-col data_unit (or None). Leaf first, then inherited
                 from the colspan'd parent above (e.g. parent <th colspan="2"
                 data-idx="BAS_001" data-unit="万元"> propagates "万元" to
                 both leaf cells in its colspan range).
    """
    n_hdr_rows = len(headers)
    n_cols = max((sum(_span(th, "colspan") for th in row) for row in headers), default=1)
    occupied = [[False] * n_cols for _ in range(n_hdr_rows)]
    spans: list[tuple[int, int, int, int, dict]] = []

    for r, row in enumerate(headers):
        c = 0
        for th in row:
            while c < n_cols and occupied[r][c]:
                c += 1
            if c >= n_cols:
                break
            rs = _span(th, "rowspan")
            cs = _span(th, "colspan")
            spans.append((r, c, rs, cs, _as_dict(th)))
            for dr in range(rs):
                for dc in range(cs):
                    rr, cc = r + dr, c + dc
                    if rr < n_hdr_rows and cc < n_cols:
                        occupied[rr][cc] = True
            c += cs

    last_r = n_hdr_rows - 1
    col_sources: list[dict | None] = [None] * n_cols
    col_units: list[str | None] = [None] * n_cols
    for col_idx in range(n_cols):
        leaf_unit = None
        # 1. Leaf first (last-row span covering this col)
        for r, c, rs, cs, th_d in spans:
            if r == last_r and c <= col_idx < c + cs:
                col_sources[col_idx] = th_d
                leaf_unit = th_d.get("data_unit")
                break
        if col_sources[col_idx] is not None:
            if leaf_unit:
                col_units[col_idx] = leaf_unit
            else:
                # 2. Inherit from any colspan'd parent above whose range covers this col
                for r, c, rs, _cs, th_d in spans:
                    if r < last_r and c <= col_idx < c + _span(th_d, "colspan") and th_d.get("data_unit"):
                        col_units[col_idx] = th_d["data_unit"]
                        break
            continue
        # 3. No leaf at this col: deepest rowspan from above owns it
        for r, c, rs, _cs, th_d in spans:
            if r < last_r and c <= col_idx < c + rs:
                col_sources[col_idx] = th_d
                col_units[col_idx] = th_d.get("data_unit")
                break
    return n_cols, spans, col_sources, col_units


def _extract_value(th_d: dict | None, row: dict) -> tuple[str, str | None]:
    """Resolve a data cell value + data_unit from its source header cell."""
    if th_d is None:
        return "", None
    if th_d.get("is_computed"):
        return row.get(th_d.get("text", ""), "—"), th_d.get("data_unit")
    if th_d.get("idx_id") and th_d.get("period"):
        return row.get(f"{th_d['idx_id']}@{th_d['period']}", "—"), th_d.get("data_unit")
    if th_d.get("text") in _ORG_LABELS:
        return row.get("branch_num", ""), None
    return "", th_d.get("data_unit")


def _render_report(docx, report: dict, style: dict) -> None:
    p = docx.add_paragraph()
    run = p.add_run(report["title"])
    _apply_font(run, style["font"]["report"])
    if report.get("description"):
        p = docx.add_paragraph()
        run = p.add_run(str(report["description"]).strip())
        _apply_font(run, style["font"]["body"])

    rows = report.get("rows", [])
    headers = report.get("headers", [])
    if not headers or not rows:
        return

    n_cols, spans, col_sources, col_units = _build_header_layout(headers)
    n_hdr_rows = len(headers)
    table = docx.add_table(rows=n_hdr_rows + len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r, c, rs, cs, _th in spans:
        if rs > 1 or cs > 1:
            table.rows[r].cells[c].merge(table.rows[r + rs - 1].cells[c + cs - 1])

    title_font = style["font"]["title"]
    body_font = style["font"]["body"]
    for r, c, _rs, _cs, th_d in spans:
        label = th_d.get("text", "")
        sub = {"text": f"({th_d['data_unit']})"} if th_d.get("data_unit") else None
        _set_cell_text(table.rows[r].cells[c], label, main_font=title_font, sub_font=sub)

    for d_idx, row in enumerate(rows):
        for col_idx, th_d in enumerate(col_sources):
            val, data_unit = _extract_value(th_d, row)
            # Prefer layout-resolved unit (parent-inherited) over leaf-direct,
            # so multi-row headers where data_unit sits on a colspan'd parent
            # cell still get correct DATA_TYPE_MAP routing.
            effective_unit = col_units[col_idx] or data_unit
            data_type = DATA_TYPE_MAP.get(effective_unit or "", "number")
            _set_cell_text(
                table.rows[n_hdr_rows + d_idx].cells[col_idx],
                _format_value(val, data_type),
                main_font=body_font,
            )


def render_docx(payload: dict, *, out_path: str, style_path: str) -> None:
    style = _load_style(style_path)
    docx = Document()

    section = docx.sections[0]
    page = style.get("page", {})
    margins = page.get("margins_cm", {})
    for k, cm in margins.items():
        setattr(section, f"{k}_margin", Cm(cm))

    p = docx.add_paragraph()
    run = p.add_run(payload["title"])
    _apply_font(run, style["font"]["title"])

    for sec in payload.get("sections", []):
        p = docx.add_paragraph()
        run = p.add_run(sec["title"])
        _apply_font(run, style["font"]["section"])
        for report in sec.get("reports", []):
            _render_report(docx, report, style)

    docx.save(out_path)