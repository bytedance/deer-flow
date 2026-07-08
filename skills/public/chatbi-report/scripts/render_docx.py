"""Render the final DOCX (`report.docx`).

依规格 §"表头副标渲染规则":
- 主列标题读取 `headers[].text`（来自 MD 的中文显示名）
  —— 不是 SQLBot idx_id，也不是 SQLBot idx_name 查询结果。
- 副标题仅为 `(data-unit)`（如 `(个)`）。
- 多级 thead 通过跨 rowspan/colspan 的 cell.merge() 渲染。
- 若 report 挂载 `description_text`，在 report heading 与 table 之间渲染描述段。
"""

from __future__ import annotations

import argparse
import json
import sys
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor

from render_markdown import attach_description_files, doc_from_dict, normalize_wide_by_report

DATA_TYPE_MAP = {
    "元": "currency",
    "万元": "currency",
    "亿元": "currency",
    "%": "percentage",
    "百分点": "ratio",
}


def _load_style(style_path: str) -> dict:
    return json.loads(Path(style_path).read_text(encoding="utf-8"))


def _apply_font(run, font_cfg: dict) -> None:
    run.font.name = font_cfg.get("name", "宋体")
    run.font.size = Pt(font_cfg.get("size", 11))
    run.font.bold = bool(font_cfg.get("bold", False))
    if "color" in font_cfg:
        run.font.color.rgb = RGBColor.from_string(font_cfg["color"].lstrip("#"))


def _set_cell_text(cell, text: str, *, main_font: dict, sub_font: dict | None = None) -> None:
    """替换单元格内容为 `text`（可选的副标在第二行）。"""
    # 清除已有段落（python-docx 单元格默认带一个空段落）
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    run = p.add_run(text)
    _apply_font(run, main_font)
    if sub_font:
        sub_p = cell.add_paragraph()
        sub_run = sub_p.add_run(sub_font["text"])
        _apply_font(sub_run, {**main_font, **sub_font})


def _format_value(value, data_type: str, style: dict) -> str:
    if value in (None, "", "⚠️QUERY_FAILED", "⚠️COMPUTE_FAILED"):
        return str(value) if value else ""
    try:
        v = float(Decimal(str(value)))
    except Exception:
        return str(value)
    if data_type == "percentage":
        # 以小数存储（0.1833）；按 1 位小数显示为百分比
        return f"{v * 100:.1f}%"
    if data_type == "currency":
        return f"¥{v:,.2f}"
    if data_type == "ratio":
        return f"{v:.2f}"
    return f"{v:,.0f}"


def _value_key(th) -> str | None:
    if getattr(th, "is_computed", False):
        return th.text.strip("{}") if th.text.startswith("{{") else th.text
    if getattr(th, "idx_id", None):
        return f"{th.idx_id}@{th.period}" if getattr(th, "period", None) else th.idx_id
    return None


def _header_grid(headers: list[list]) -> list[list]:
    grid: list[list] = []
    for r_idx, row in enumerate(headers):
        while len(grid) <= r_idx:
            grid.append([])
        c_idx = 0
        for th in row:
            while c_idx < len(grid[r_idx]) and grid[r_idx][c_idx] is not None:
                c_idx += 1
            rowspan = th.rowspan or 1
            colspan = th.colspan or 1
            for rr in range(r_idx, r_idx + rowspan):
                while len(grid) <= rr:
                    grid.append([])
                while len(grid[rr]) < c_idx + colspan:
                    grid[rr].append(None)
            for rr in range(r_idx, r_idx + rowspan):
                for cc in range(c_idx, c_idx + colspan):
                    grid[rr][cc] = th
            c_idx += colspan
    width = max((len(row) for row in grid), default=0)
    for row in grid:
        while len(row) < width:
            row.append(None)
    return grid


def _leaf_cells(headers: list[list]) -> list:
    grid = _header_grid(headers)
    return [c for c in grid[-1] if c is not None] if grid else []


def _add_styled_paragraph(docx, text: str, font_cfg: dict) -> None:
    p = docx.add_paragraph()
    run = p.add_run(text)
    _apply_font(run, font_cfg)


def _resolve_chart_png(
    chart_entry: dict, *, chart_dir: str | Path | None,
) -> Path | None:
    """Resolve a chart_entry's PNG path. Returns None if path unsafe / missing.

    Defect 2 fix: prefer relative_path + chart_dir for cross-process safety
    (manifest can move between chart_gen and render_docx); fall back to the
    absolute path only when chart_dir is unknown or relative_path is missing.

    Path-traversal check: when chart_dir is given, the resolved path must
    live inside it. Reject any PNG that escapes (e.g., manifest corruption,
    user-controlled relative_path).
    """
    rel = chart_entry.get("relative_path", "")
    abs_path_str = chart_entry.get("path", "")

    candidate: Path | None = None
    if chart_dir and rel:
        # rel is "{stem}.charts/{slug}.png"; chart_dir is "{out_dir}/{stem}.charts".
        # The PNG lives at chart_dir / basename(rel).
        candidate = (Path(chart_dir) / Path(rel).name).resolve()
    elif abs_path_str:
        candidate = Path(abs_path_str).resolve()
    else:
        return None

    if chart_dir:
        try:
            candidate.relative_to(Path(chart_dir).resolve())
        except ValueError:
            print(
                f"WARN: chart PNG escapes chart_dir: {candidate} not under {chart_dir}",
                file=sys.stderr,
            )
            return None
    return candidate


def _build_chart_lookup(chart_manifest: dict | None) -> dict[tuple[int, int], list[dict]]:
    """Index chart_manifest.reports by (section_idx, report_idx) for O(1) lookup."""
    if not chart_manifest:
        return {}
    lookup: dict[tuple[int, int], list[dict]] = {}
    for entry in chart_manifest.get("reports", []) or []:
        key = (int(entry.get("section_idx", -1)), int(entry.get("report_idx", -1)))
        if key[0] < 0 or key[1] < 0:
            continue
        lookup[key] = list(entry.get("charts", []) or [])
    return lookup


def _embed_chart(
    docx, chart_entry: dict, style: dict,
    *, chart_dir: str | Path | None = None,
    out_path: str | None = None,
) -> None:
    """Embed a single chart PNG into the DOCX under the report's data table.

    `chart_entry` is the manifest dict produced by chart_gen.generate_charts:
      {title, type, status, path, relative_path, ...}.

    Failures (missing file, bad status, path escapes chart_dir) are logged
    to stderr; the renderer keeps going so the report's table still ships
    even when the chart fails.
    """
    status = chart_entry.get("status")
    if status != "ok":
        print(
            f"WARN: skipping chart `{chart_entry.get('title', '')}` "
            f"(status={status}, error={chart_entry.get('error', '')})",
            file=sys.stderr,
        )
        return
    png_path = _resolve_chart_png(chart_entry, chart_dir=chart_dir)
    if png_path is None:
        print(
            f"WARN: chart PNG cannot be resolved (manifest corrupt or escapes chart_dir)",
            file=sys.stderr,
        )
        return
    if not png_path.exists():
        print(
            f"WARN: chart PNG missing on disk: {png_path}",
            file=sys.stderr,
        )
        return
    title = chart_entry.get("title") or ""
    if title:
        _add_styled_paragraph(docx, title, style["font"]["report"])
    p = docx.add_paragraph()
    run = p.add_run()
    run.add_picture(str(png_path), width=Cm(14))
    print(
        f"OK: embedded chart `{title}` -> {out_path or '?'}",
        file=sys.stderr,
    )


def render_docx(
    report_doc,
    wide_by_report: list[list[dict]],
    *,
    out_path: str,
    style_path: str,
    chart_manifest: dict | None = None,
    chart_dir: str | Path | None = None,
) -> None:
    """渲染完整 DOCX。"""
    style = _load_style(style_path)
    docx = Document()

    # 页面设置
    section = docx.sections[0]
    page = style.get("page", {})
    margins = page.get("margins_cm", {})
    if page.get("orientation") == "landscape":
        from docx.enum.section import WD_ORIENTATION

        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for k, cm in margins.items():
        setattr(section, f"{k}_margin", Cm(cm))

    # 标题
    p = docx.add_paragraph()
    run = p.add_run(report_doc.title)
    _apply_font(run, style["font"]["title"])

    # wide_by_report is positionally aligned with report iteration order
    # (Orchestrator._finish_phase_2 builds it that way); track cumulative
    # offset as we walk sections so each report slices its own row block.
    chart_lookup = _build_chart_lookup(chart_manifest)
    wide_offset = 0
    for sec_idx, sec in enumerate(report_doc.sections):
        _render_section(
            docx, sec, wide_by_report,
            sec_idx=sec_idx,
            wide_offset=wide_offset,
            style=style,
            chart_lookup=chart_lookup,
            chart_dir=chart_dir,
            out_path=out_path,
        )
        wide_offset += len(sec.reports)
    docx.save(out_path)


def _render_section(
    docx, sec, wide_by_report, *, sec_idx: int, wide_offset: int, style: dict,
    chart_lookup: dict[tuple[int, int], list[dict]] | None = None,
    chart_dir: str | Path | None = None,
    out_path: str | None = None,
):
    p = docx.add_paragraph()
    run = p.add_run(sec.title)
    _apply_font(run, style["font"]["section"])

    for rep_idx, report in enumerate(sec.reports):
        pos = wide_offset + rep_idx
        rows = wide_by_report[pos] if pos < len(wide_by_report) else []
        charts = (chart_lookup or {}).get((sec_idx, rep_idx), [])
        _render_report(
            docx, report, rows, style,
            charts=charts, chart_dir=chart_dir, out_path=out_path,
        )


def _render_report(
    docx, report, wide_rows, style,
    *, charts: list[dict] | None = None,
    chart_dir: str | Path | None = None,
    out_path: str | None = None,
):
    p = docx.add_paragraph()
    run = p.add_run(report.title)
    _apply_font(run, style["font"]["report"])

    description_text = getattr(report, "description_text", None)
    if description_text:
        _add_styled_paragraph(docx, str(description_text).strip(), style["font"]["body"])

    # Embed charts AFTER description + BEFORE data table — readers understand
    # numbers in context of the visualization they explain.
    if charts:
        for chart_entry in charts:
            _embed_chart(
                docx, chart_entry, style,
                chart_dir=chart_dir, out_path=out_path,
            )

    if not wide_rows:
        docx.add_paragraph().add_run("（无数据行）").italic = True
        return

    leaves = _leaf_cells(report.headers)
    n_cols = len(leaves) or max((len(row) for row in report.headers), default=0) or 1
    n_rows = len(report.headers) + len(wide_rows)
    table = docx.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # 表头行
    for r_idx, header_row in enumerate(report.headers):
        for c_idx, cell_def in enumerate(header_row):
            if c_idx >= n_cols:
                break
            tc = table.rows[r_idx].cells[c_idx]
            label = cell_def.text or ""
            sub = None
            if cell_def.data_unit:
                sub = {"text": f"({cell_def.data_unit})"}
            _set_cell_text(tc, label, main_font=style["font"]["title" if r_idx == 0 else "section"], sub_font=sub)
            # 背景
            tc._tc.get_or_add_tcPr()

    # 数据行：按叶子表头网格对齐，保留 rowspan 维度列
    for d_idx, row in enumerate(wide_rows):
        cells = row.get("cells", {})
        for c_idx, th in enumerate(leaves):
            if c_idx >= n_cols:
                break
            tc = table.rows[len(report.headers) + d_idx].cells[c_idx]
            key = _value_key(th)
            if key:
                val = cells.get(key, "—")
                data_type = DATA_TYPE_MAP.get(th.data_unit or "", "number")
                text = _format_value(val, data_type, style)
            else:
                label = (th.text or "").strip()
                if label in {"季度", "时期", "日期", "period"}:
                    text = str(row.get("data_dt", ""))
                elif label in {"机构", "行社", "分行", "网点", "org"}:
                    text = str(row.get("org_ecd", ""))
                else:
                    text = ""
            _set_cell_text(tc, text, main_font=style["font"]["body"])


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="render_docx", description=__doc__.splitlines()[0])
    parser.add_argument("--parsed", required=True)
    parser.add_argument("--wide", required=True)
    parser.add_argument("--style", required=True)
    parser.add_argument("--out", required=True)
    parser.add_argument("--descriptions-dir", default=None)
    parser.add_argument("--stem", default=None)
    parser.add_argument(
        "--chart-manifest", default=None,
        help="chart_gen manifest JSON (optional; charts are skipped when absent)",
    )
    parser.add_argument(
        "--chart-dir", default=None,
        help="directory containing chart PNGs; defaults to <stem>.charts/ "
        "sibling of the manifest when both are provided (Defect 2: cross-process safe)",
    )
    args = parser.parse_args(argv)

    try:
        parsed = json.loads(Path(args.parsed).read_text(encoding="utf-8"))
        doc = doc_from_dict(parsed)
        wide = normalize_wide_by_report(doc, json.loads(Path(args.wide).read_text(encoding="utf-8")))
        attach_description_files(doc, args.descriptions_dir, args.stem or Path(args.parsed).name.removesuffix(".parsed.json"))
        chart_manifest = None
        chart_dir = args.chart_dir
        if args.chart_manifest:
            cm_path = Path(args.chart_manifest)
            if cm_path.exists():
                chart_manifest = json.loads(cm_path.read_text(encoding="utf-8"))
                # Defect 2 fallback: derive chart_dir from stem next to manifest
                # when caller did not pass --chart-dir explicitly.
                if chart_dir is None:
                    stem = args.stem or cm_path.name.removesuffix(".charts.json")
                    chart_dir = str(cm_path.parent / f"{stem}.charts")
            else:
                print(
                    f"WARN: --chart-manifest specified but file missing: {cm_path}; charts skipped",
                    file=sys.stderr,
                )
        render_docx(
            doc, wide, out_path=args.out, style_path=args.style,
            chart_manifest=chart_manifest, chart_dir=chart_dir,
        )
    except Exception as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        return 1
    print(f"OK: rendered docx -> {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
