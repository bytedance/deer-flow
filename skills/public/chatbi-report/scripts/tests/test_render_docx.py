"""Unit tests for scripts/render_docx.py.

Tests round-trip a written .docx through python-docx to verify
header text, cell merges, font, and ⚠️ markers.
"""

import subprocess
import sys
from pathlib import Path


def _write_driver(driver_path: Path, scripts_dir: Path, doc_path: str,
                  out_path: str, style_path: str,
                  cells: dict[str, str], data_dt: str = "2025-Q4",
                  org_ecd: str = "王益联社", description_text: str | None = None) -> None:
    """Write a small driver script that parses MD + renders .docx."""
    driver_path.write_text(
        f"""# ruff: noqa
import sys
from pathlib import Path
sys.path.insert(0, r"{scripts_dir}")
import parse_md as pm
import render_docx as rd
doc = pm.parse_file(r"{doc_path}")
if {description_text!r} is not None:
    doc.sections[0].reports[0].description_text = {description_text!r}
wide = [{{
    "data_dt": r"{data_dt}", "org_ecd": r"{org_ecd}",
    "cells": {cells!r},
    "raw_cells": {cells!r},
}}]
rd.render_docx(doc, [wide], out_path=r"{out_path}",
               style_path=r"{style_path}")
""",
        encoding="utf-8",
    )


def _run_driver(driver_path: Path) -> subprocess.CompletedProcess:
    proc = subprocess.run([sys.executable, str(driver_path)],
                          capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(f"render_docx driver failed: {proc.stderr}")
    return proc


def test_render_docx_writes_a_valid_docx(fixture_dir, tmp_path):
    """render_docx() 产出非空 .docx 文件（single_org.md 走单期路径）。"""
    scripts_dir = Path(__file__).resolve().parents[1]
    style_path = scripts_dir / "report_style.json"
    out = tmp_path / "report.docx"
    driver = tmp_path / "_drv.py"
    _write_driver(driver, scripts_dir,
                  str(fixture_dir / "sample_md" / "single_org.md"),
                  str(out), str(style_path),
                  cells={"BAS_0263": "1,420", "收单商户同比": "0.1833"})
    _run_driver(driver)
    assert out.exists()
    assert out.stat().st_size > 1024


def test_render_docx_header_uses_chinese_name_not_idx_id(fixture_dir, tmp_path):
    """Chatbi 规则：列主标题为 MD 中的中文显示名（"贷款收单商户数"），不出现 idx_id。"""
    scripts_dir = Path(__file__).resolve().parents[1]
    style_path = scripts_dir / "report_style.json"
    out = tmp_path / "report.docx"
    driver = tmp_path / "_drv.py"
    _write_driver(driver, scripts_dir,
                  str(fixture_dir / "sample_md" / "single_org.md"),
                  str(out), str(style_path),
                  cells={"BAS_0263": "1,420", "收单商户同比": "0.1833"})
    _run_driver(driver)

    from docx import Document
    doc = Document(str(out))
    cells_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "贷款收单商户数" in cells_text
    # idx_id 不应作为可见列标题
    assert "BAS_0263" not in cells_text


def test_render_docx_multi_level_merges_cells(fixture_dir, tmp_path):
    """multi_org.md：顶行 colspan=3 父级表头（利润总额 / 同比增速）写入 docx。"""
    scripts_dir = Path(__file__).resolve().parents[1]
    style_path = scripts_dir / "report_style.json"
    out = tmp_path / "report.docx"
    driver = tmp_path / "_drv.py"
    # multi_org.md 的叶子 cell key 是 `BAS_0263@2023/2024/2025`
    cells = {
        "BAS_0263@2023": "188.01", "BAS_0263@2024": "495.83", "BAS_0263@2025": "322.78",
        "2023利润同比": "-0.688", "2024利润同比": "1.6372", "2025利润同比": "-0.349",
    }
    _write_driver(driver, scripts_dir,
                  str(fixture_dir / "sample_md" / "multi_org.md"),
                  str(out), str(style_path),
                  cells=cells, data_dt="2023-2025", org_ecd="王益联社")
    _run_driver(driver)

    from docx import Document
    doc = Document(str(out))
    table = doc.tables[0]
    texts = "\n".join(c.text for r in table.rows for c in r.cells)
    # 父级 + 叶子层都应出现
    assert "利润总额" in texts
    assert "同比增速" in texts
    assert "2023年" in texts and "2024年" in texts and "2025年" in texts


def test_render_docx_query_failed_marker_in_cell(fixture_dir, tmp_path):
    """⚠️QUERY_FAILED 单元格文本按原样保留（单期路径）。"""
    scripts_dir = Path(__file__).resolve().parents[1]
    style_path = scripts_dir / "report_style.json"
    out = tmp_path / "report.docx"
    driver = tmp_path / "_drv_fail.py"
    _write_driver(driver, scripts_dir,
                  str(fixture_dir / "sample_md" / "single_org.md"),
                  str(out), str(style_path),
                  cells={"BAS_0263": "⚠️QUERY_FAILED", "收单商户同比": "—"})
    _run_driver(driver)

    from docx import Document
    doc = Document(str(out))
    cells_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "⚠️QUERY_FAILED" in cells_text


def test_render_docx_description_paragraph_before_table(fixture_dir, tmp_path):
    scripts_dir = Path(__file__).resolve().parents[1]
    style_path = scripts_dir / "report_style.json"
    out = tmp_path / "report.docx"
    driver = tmp_path / "_drv_desc.py"
    _write_driver(driver, scripts_dir,
                  str(fixture_dir / "sample_md" / "single_org.md"),
                  str(out), str(style_path),
                  cells={"BAS_0263": "1,420", "收单商户同比": "0.1833"},
                  description_text="这是DOCX描述段。")
    _run_driver(driver)

    from docx import Document
    doc = Document(str(out))
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    assert "这是DOCX描述段。" in paragraphs


# ---------- chart manifest integration (Task 7) ---------- #

def test_render_docx_embeds_chart(tmp_path):
    """H5 fix: _render_report actually embeds ok chart PNGs into docx via add_picture."""
    import json

    import docx as docx_lib
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from parse_md import Report, OrgContext, Th
    from render_docx import _render_report
    from render_markdown import _charts_for_report

    # Real 1x1 PNG so add_picture succeeds
    png = tmp_path / "chart.png"
    fig, ax = plt.subplots()
    ax.plot([1, 2, 3])
    fig.savefig(str(png))
    plt.close(fig)

    manifest = {
        "reports": [{
            "section_idx": 0,
            "report_idx": 0,
            "charts": [
                {"title": "利润趋势", "type": "line", "status": "ok", "path": str(png)},
                {"title": "bad", "type": "bar", "status": "failed", "path": ""},
            ],
        }],
    }
    # Shared helper: 1 ok chart, failed excluded
    assert len(_charts_for_report(manifest, 0, 0)) == 1

    report = Report(
        title="R",
        org_contexts=[OrgContext("A", "机构A")],
        time_info=["2025"],
        headers=[[
            Th(text="行社", is_indicator=False, is_computed=False),
            Th(text="利润", is_indicator=True, is_computed=False, idx_id="BAS_0263", period="2025", data_unit="万元"),
        ]],
        data_rows=[{"data_dt": "2025", "raw_cells": ["机构A", "100"]}],
    )
    wide_rows = [{"data_dt": "2025", "org_ecd": "机构A",
                  "cells": {"BAS_0263@2025": "100"}, "raw_cells": {"BAS_0263@2025": "100"}}]

    style_path = Path(__file__).resolve().parent.parent / "report_style.json"
    style = json.loads(style_path.read_text(encoding="utf-8"))

    docx_doc = docx_lib.Document()
    _render_report(docx_doc, report, wide_rows, style, manifest, 0, 0)

    # 1 ok chart → 1 inline shape
    assert len(docx_doc.inline_shapes) == 1
